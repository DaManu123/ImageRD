"""
exporters.py — Exportadores de resultados OCR.

Genera archivos de salida en múltiples formatos:
  - Texto plano (.txt)
  - JSON estructurado (.json)
  - Documento Word (.docx)

Cada exportador recibe el resultado OCR y el documento reconstruido,
y genera un archivo limpio y profesional.
"""

import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from layout_reconstructor import ReconstructedDocument, analyze_document_structure
from ocr_engine import BlockData, OCRResult
from utils import generate_output_path, logger


# ─────────────────────────────────────────────
# Exportador a texto plano
# ─────────────────────────────────────────────
def export_to_txt(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
    output_path: Optional[str] = None,
    input_path: Optional[str] = None,
) -> Path:
    """
    Exporta el texto extraído a un archivo .txt limpio.

    El formato preserva:
      - Saltos de línea originales.
      - Separación entre bloques.
      - Indentación detectada.

    Args:
        ocr_result: Resultado de la extracción OCR.
        reconstructed: Documento reconstruido con estructura.
        output_path: Ruta de salida personalizada (opcional).
        input_path: Ruta de la imagen de entrada (para generar nombre).

    Returns:
        Path del archivo generado.
    """
    if output_path:
        path = Path(output_path)
    elif input_path:
        path = generate_output_path(Path(input_path), "txt")
    else:
        path = Path("resultado.txt")

    logger.info("Exportando a TXT: %s", path)

    content = _build_txt_content(ocr_result, reconstructed)

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")

    logger.info("Archivo TXT generado correctamente (%d bytes).", path.stat().st_size)
    return path


def _build_txt_content(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
) -> str:
    """
    Construye el contenido del archivo TXT.

    Args:
        ocr_result: Resultado OCR.
        reconstructed: Documento reconstruido.

    Returns:
        Contenido formateado para el archivo.
    """
    lines = [
        "=" * 60,
        "TEXTO EXTRAÍDO POR OCR",
        "=" * 60,
        "",
        f"Idioma: {ocr_result.language}",
        f"Palabras detectadas: {ocr_result.word_count}",
        f"Confianza promedio: {ocr_result.avg_confidence:.1f}%",
        f"Bloques: {len(ocr_result.blocks)}",
        f"Columnas: {reconstructed.column_count}",
        "",
        "-" * 60,
        "",
        reconstructed.formatted_text,
        "",
        "-" * 60,
        "Fin del documento",
    ]

    return "\n".join(lines)


# ─────────────────────────────────────────────
# Exportador a JSON estructurado
# ─────────────────────────────────────────────
def export_to_json(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
    output_path: Optional[str] = None,
    input_path: Optional[str] = None,
    include_words: bool = False,
) -> Path:
    """
    Exporta los datos OCR a un archivo JSON estructurado.

    Incluye:
      - Metadatos del documento.
      - Bloques con coordenadas.
      - Texto por bloque y línea.
      - Estructura de columnas.

    Args:
        ocr_result: Resultado de la extracción OCR.
        reconstructed: Documento reconstruido.
        output_path: Ruta de salida personalizada.
        input_path: Ruta de la imagen de entrada.
        include_words: Si incluir datos individuales de cada palabra.

    Returns:
        Path del archivo generado.
    """
    if output_path:
        path = Path(output_path)
    elif input_path:
        path = generate_output_path(Path(input_path), "json")
    else:
        path = Path("resultado.json")

    logger.info("Exportando a JSON: %s", path)

    data = _build_json_structure(ocr_result, reconstructed, include_words)

    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    logger.info("Archivo JSON generado correctamente.")
    return path


def _build_json_structure(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
    include_words: bool = False,
) -> Dict[str, Any]:
    """
    Construye la estructura JSON del documento.

    Args:
        ocr_result: Resultado OCR.
        reconstructed: Documento reconstruido.
        include_words: Incluir palabras individuales.

    Returns:
        Diccionario listo para serializar a JSON.
    """
    # Metadatos
    metadata = {
        "idioma": ocr_result.language,
        "total_palabras": ocr_result.word_count,
        "confianza_promedio": round(ocr_result.avg_confidence, 2),
        "dimensiones_imagen": {
            "ancho": ocr_result.image_width,
            "alto": ocr_result.image_height,
        },
        "es_multicolumna": reconstructed.is_multicolumn,
        "total_columnas": reconstructed.column_count,
        "total_bloques": len(ocr_result.blocks),
    }

    # Bloques
    bloques = []
    for block in reconstructed.ordered_blocks:
        bloque_data = _serialize_block(block, include_words)
        bloques.append(bloque_data)

    # Columnas
    columnas = []
    for col in reconstructed.columns:
        columnas.append({
            "columna_id": col.column_id,
            "x_inicio": col.x_start,
            "x_fin": col.x_end,
            "bloques_ids": [b.block_num for b in col.blocks],
        })

    return {
        "metadata": metadata,
        "texto_completo": reconstructed.formatted_text,
        "bloques": bloques,
        "columnas": columnas,
    }


def _serialize_block(block: BlockData, include_words: bool) -> Dict[str, Any]:
    """
    Serializa un bloque para formato JSON.

    Args:
        block: Bloque de texto.
        include_words: Incluir palabras individuales.

    Returns:
        Diccionario con datos del bloque.
    """
    lineas = []
    for line in block.lines:
        linea_data: Dict[str, Any] = {
            "texto": line.text,
            "coordenadas": [line.x, line.y, line.width, line.height],
        }

        if include_words:
            linea_data["palabras"] = [
                {
                    "texto": w.text,
                    "confianza": round(w.confidence, 1),
                    "coordenadas": [w.x, w.y, w.width, w.height],
                }
                for w in line.words
            ]

        lineas.append(linea_data)

    return {
        "tipo": block.block_type,
        "bloque_num": block.block_num,
        "coordenadas": [block.x, block.y, block.width, block.height],
        "texto": "\n".join(line.text for line in block.lines),
        "lineas": lineas,
    }


# ─────────────────────────────────────────────
# Exportador a Word (.docx)
# ─────────────────────────────────────────────
def export_to_docx(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
    output_path: Optional[str] = None,
    input_path: Optional[str] = None,
) -> Path:
    """
    Exporta el texto extraído a un documento Word (.docx).

    Preserva:
      - Párrafos separados por bloque.
      - Formato básico (fuente, tamaño).
      - Encabezado con metadatos.

    Args:
        ocr_result: Resultado de la extracción OCR.
        reconstructed: Documento reconstruido.
        output_path: Ruta de salida personalizada.
        input_path: Ruta de la imagen de entrada.

    Returns:
        Path del archivo generado.
    """
    if output_path:
        path = Path(output_path)
    elif input_path:
        path = generate_output_path(Path(input_path), "docx")
    else:
        path = Path("resultado.docx")

    logger.info("Exportando a DOCX: %s", path)

    doc = DocxDocument()
    _build_docx_content(doc, ocr_result, reconstructed)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))

    logger.info("Archivo DOCX generado correctamente.")
    return path


def _build_docx_content(
    doc: DocxDocumentType,
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
) -> None:
    """
    Construye el contenido del documento Word.

    Args:
        doc: Documento Word en construcción.
        ocr_result: Resultado OCR.
        reconstructed: Documento reconstruido.
    """
    # ── Título ──
    title = doc.add_heading("Texto Extraído por OCR", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadatos ──
    analysis = analyze_document_structure(ocr_result)
    meta_paragraph = doc.add_paragraph()
    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_lines = [
        f"Idioma: {ocr_result.language}",
        f"Palabras: {ocr_result.word_count}",
        f"Confianza: {ocr_result.avg_confidence:.1f}%",
        f"Bloques: {len(ocr_result.blocks)}",
        f"Columnas: {reconstructed.column_count}",
    ]

    for line in meta_lines:
        run = meta_paragraph.add_run(line + "\n")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("─" * 50)

    # ── Contenido por bloques ──
    for i, block in enumerate(reconstructed.ordered_blocks):
        if not block.lines:
            continue

        # Cada bloque es un párrafo
        for line in block.lines:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"

        # Separador entre bloques (excepto el último)
        if i < len(reconstructed.ordered_blocks) - 1:
            doc.add_paragraph("")

    # ── Pie de página ──
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph("Generado por ImageRD — Sistema OCR Profesional")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(160, 160, 160)


# ─────────────────────────────────────────────
# Exportador maestro
# ─────────────────────────────────────────────
def export_all(
    ocr_result: OCRResult,
    reconstructed: ReconstructedDocument,
    output_format: str = "txt",
    output_path: Optional[str] = None,
    input_path: Optional[str] = None,
) -> List[Path]:
    """
    Exporta el resultado OCR al formato o formatos solicitados.

    Args:
        ocr_result: Resultado de la extracción OCR.
        reconstructed: Documento reconstruido.
        output_format: Formato deseado ('txt', 'json', 'docx', 'all').
        output_path: Ruta de salida (solo para formato único).
        input_path: Ruta de la imagen de entrada.

    Returns:
        Lista de Paths de los archivos generados.
    """
    generated: List[Path] = []
    fmt = output_format.lower().strip()

    if fmt in ("txt", "all"):
        path = export_to_txt(
            ocr_result, reconstructed,
            output_path=output_path if fmt == "txt" else None,
            input_path=input_path,
        )
        generated.append(path)

    if fmt in ("json", "all"):
        path = export_to_json(
            ocr_result, reconstructed,
            output_path=output_path if fmt == "json" else None,
            input_path=input_path,
            include_words=True,
        )
        generated.append(path)

    if fmt in ("docx", "all"):
        path = export_to_docx(
            ocr_result, reconstructed,
            output_path=output_path if fmt == "docx" else None,
            input_path=input_path,
        )
        generated.append(path)

    return generated
